#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
公文自动格式化工具（GB/T 9704-2012）
作者：茗记
版本：4.0.0

改进说明：
- 落款支持手动偏移量校正，解决字体宽度差异导致的居中偏差
- 所有段落强制清除首行缩进干扰
- 落款数字与汉字统一使用仿宋体
- v4.0.0: 全面代码质量优化（类型注解、错误处理、性能提升、安全性增强）
"""

import sys
import os
import re
import shutil
import logging
from functools import lru_cache
from typing import Optional, List
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
from docx.text.paragraph import Paragraph
from docx.document import Document as DocumentType
from docx.styles.style import _ParagraphStyle

# ------------------------------ 可调参数 ---------------------------------------
BODY_FONT = '仿宋_GB2312'
BODY_SIZE = 16
PAGE_FORMAT = '— PAGE —'
SIGNATURE_OFFSET = 0.0  # 单位左移校正值（正数左移，负数右移），单位：全角字符
# ----------------------------------------------------------------------------

# ------------------------------ 样式配置常量 -----------------------------------
STYLE_CONFIG = {
    '公文标题': {
        'font': '方正小标宋简体',
        'size': 22,
        'bold': False,
        'alignment': WD_ALIGN_PARAGRAPH.CENTER,
        'first_line_indent': None,
        'space_before': 0,
        'space_after': 0,
        'line_spacing': 32,
        'set_west': True,
    },
    '一级标题': {
        'font': '黑体',
        'size': BODY_SIZE,
        'bold': False,
        'alignment': None,
        'first_line_indent': 0,
        'space_before': 0,
        'space_after': 0,
        'line_spacing': 28,
        'set_west': True,
    },
    '二级标题': {
        'font': '楷体_GB2312',
        'size': BODY_SIZE,
        'bold': True,
        'alignment': None,
        'first_line_indent': 0,
        'space_before': 0,
        'space_after': 0,
        'line_spacing': 28,
        'set_west': True,
    },
    '主送机关': {
        'font': BODY_FONT,
        'size': BODY_SIZE,
        'bold': False,
        'alignment': WD_ALIGN_PARAGRAPH.LEFT,
        'first_line_indent': 0,
        'space_before': 0,
        'space_after': 0,
        'line_spacing': 28,
        'set_west': True,
    },
    '公文正文': {
        'font': BODY_FONT,
        'size': BODY_SIZE,
        'bold': False,
        'alignment': None,
        'first_line_indent': 2,
        'space_before': 0,
        'space_after': 0,
        'line_spacing': 28,
        'set_west': True,
    },
    '公文落款': {
        'font': BODY_FONT,
        'size': BODY_SIZE,
        'bold': False,
        'alignment': WD_ALIGN_PARAGRAPH.RIGHT,
        'first_line_indent': 0,
        'right_indent': 0,
        'space_before': 0,
        'space_after': 0,
        'line_spacing': 28,
        'set_west': False,
    },
    '公文附件': {
        'font': BODY_FONT,
        'size': BODY_SIZE,
        'bold': False,
        'alignment': None,
        'first_line_indent': 0,
        'space_before': 0,
        'space_after': 0,
        'line_spacing': 28,
        'set_west': True,
    },
}

# ------------------------------ 页面配置常量 -----------------------------------
PAGE_HEIGHT = Cm(29.7)
PAGE_WIDTH = Cm(21.0)
TOP_MARGIN = Cm(3.7)
BOTTOM_MARGIN = Cm(3.5)
LEFT_MARGIN = Cm(2.8)
RIGHT_MARGIN = Cm(2.6)

# ------------------------------ 行距配置常量 -----------------------------------
LINE_SPACING_TITLE = 32
LINE_SPACING_OTHER = 28
FOOTER_FONT_SIZE = 14

# ------------------------------ 正则表达式预编译 -------------------------------
RE_H1 = re.compile(r'^[一二三四五六七八九十]+、')
RE_H2 = re.compile(r'^（[一二三四五六七八九十]+）')
RE_ATTACH = re.compile(r'^附件\d*\s*[：:]')
RE_MAIN_RECEIVER = re.compile(
    r'^(?!.*(?:注明|备注|说明|解释|参考|来源|联系|电话|传真|电子邮箱|附件|抄送|主题词|印发日期))[^。；，！？…]+[：:]\s*$'
)
RE_DATE_PATTERN = re.compile(
    r'(\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日)'
    r'|((?:二[零〇一二三四五六七八九十]+)\s*年\s*.*?\s*月\s*.*?\s*日)'
)
RE_ATTACH_PREFIX = re.compile(r'^(附件\d*\s*[：:])')
RE_ATTACH_INDENT = re.compile(r'^\d+[.、]')
RE_SIGNATURE_EXCLUDE = re.compile(r'附件|抄送|主题词|（此页无正文）|\d{4}年')
RE_H2_SPLIT = re.compile(r'[。；]')
RE_CLEAN_SPACES = re.compile(r'[\s\u3000]')

# ------------------------------ 日志配置 ---------------------------------------
logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')
logger = logging.getLogger(__name__)


@lru_cache(maxsize=512)
def is_fullwidth(ch: str) -> bool:
    """判断字符是否为全角字符。
    
    Args:
        ch: 单个字符
        
    Returns:
        是否为全角字符
    """
    code = ord(ch)
    if 0x4E00 <= code <= 0x9FFF:
        return True
    elif 0x3400 <= code <= 0x4DBF:
        return True
    elif 0xFF00 < code <= 0xFFEF:
        return True
    elif code in (0x3000, 0x3001, 0x3002, 0x300A, 0x300B, 0x300C, 0x300D,
                  0x300E, 0x300F, 0x3010, 0x3011, 0x2018, 0x2019,
                  0x201C, 0x201D, 0x2013, 0x2014, 0xFF08, 0xFF09):
        return True
    elif 0xFE30 <= code <= 0xFE4F:
        return True
    else:
        return False


def char_width(ch: str) -> float:
    """计算字符宽度（以全角字符为单位）。
    
    Args:
        ch: 单个字符
        
    Returns:
        字符宽度，全角字符为1.0，ASCII为0.5
    """
    if is_fullwidth(ch):
        return 1.0
    if ch.isascii():
        return 0.5
    return 1.0


def get_or_create_style(doc: DocumentType, name: str, style_type=WD_STYLE_TYPE.PARAGRAPH) -> _ParagraphStyle:
    """获取或创建文档样式。
    
    Args:
        doc: Document对象
        name: 样式名称
        style_type: 样式类型
        
    Returns:
        样式对象
    """
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, style_type)


def apply_style_config(style: _ParagraphStyle, config: dict) -> None:
    """根据配置字典应用样式设置。
    
    Args:
        style: 样式对象
        config: 样式配置字典
    """
    style.font.name = config['font']
    style.font.size = Pt(config['size'])
    style.font.bold = config['bold']
    
    if config['alignment'] is not None:
        style.paragraph_format.alignment = config['alignment']
    if config['first_line_indent'] is not None:
        if config['first_line_indent'] == 0:
            style.paragraph_format.first_line_indent = Pt(0)
        else:
            style.paragraph_format.first_line_indent = pt_from_chars(BODY_SIZE, config['first_line_indent'])
    if 'right_indent' in config:
        style.paragraph_format.right_indent = Pt(config['right_indent'])
    
    style.paragraph_format.space_before = Pt(config['space_before'])
    style.paragraph_format.space_after = Pt(config['space_after'])
    style.paragraph_format.line_spacing = Pt(config['line_spacing'])
    
    rPr = style.element.get_or_add_rPr()
    safe_set_font(rPr, config['font'], config.get('set_west', True))


def pt_from_chars(font_size_pt: float, chars: float = 2.0) -> Pt:
    """将字符数转换为磅值。
    
    Args:
        font_size_pt: 字体大小（磅）
        chars: 字符数
        
    Returns:
        对应的磅值
    """
    return Pt(font_size_pt * chars)


def safe_set_font(rPr, font_name: str, set_west: bool = True) -> None:
    """安全设置字体，确保XML元素存在。
    
    Args:
        rPr: 运行属性XML元素
        font_name: 字体名称
        set_west: 是否设置西文字体
    """
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    if set_west:
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    else:
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)


def set_run_font(run, font_name: str, font_size: float, bold: bool = False, set_west: bool = True) -> None:
    """设置文本运行的字体。
    
    Args:
        run: 运行对象
        font_name: 字体名称
        font_size: 字体大小（磅）
        bold: 是否加粗
        set_west: 是否设置西文字体
    """
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    safe_set_font(rPr, font_name, set_west)


def set_paragraph_font(paragraph: Paragraph, font_name: str, font_size: float, bold: bool = False) -> None:
    """设置段落中所有运行的字体。
    
    Args:
        paragraph: 段落对象
        font_name: 字体名称
        font_size: 字体大小（磅）
        bold: 是否加粗
    """
    for run in paragraph.runs:
        set_run_font(run, font_name, font_size, bold)


def clear_run_font(run) -> None:
    """清除运行的字体格式。
    
    Args:
        run: 运行对象
    """
    run.font.name = None
    run.font.size = None
    run.font.bold = None
    run.font.italic = None
    rPr = run._element.find(qn('w:rPr'))
    if rPr is not None:
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is not None:
            for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
                if qn(attr) in rFonts.attrib:
                    del rFonts.attrib[qn(attr)]
            if not rFonts.attrib:
                rPr.remove(rFonts)


def strip_leading_spaces(paragraph: Paragraph) -> None:
    """清除段落开头的所有空白字符。
    
    Args:
        paragraph: 段落对象
    """
    runs = paragraph.runs
    if not runs:
        return
    full_text = ''.join(run.text for run in runs)
    stripped = full_text.lstrip()
    if stripped == full_text:
        return
    for run in runs:
        run._element.getparent().remove(run._element)
    new_run = paragraph.add_run(stripped)
    set_run_font(new_run, BODY_FONT, BODY_SIZE)


def insert_empty_paragraph_before(target_paragraph: Paragraph, doc: DocumentType, body_style) -> None:
    """在目标段落前插入空段落。
    
    Args:
        target_paragraph: 目标段落
        doc: Document对象
        body_style: 正文样式
    """
    empty_p = OxmlElement('w:p')
    target_paragraph._element.addprevious(empty_p)
    for p in doc.paragraphs:
        if p._element is empty_p:
            p.style = body_style
            break


def find_title_index(paragraphs: List[Paragraph]) -> Optional[int]:
    """查找标题段落的索引（单次遍历优化）。
    
    Args:
        paragraphs: 段落列表
        
    Returns:
        标题段落索引，未找到返回None
    """
    first_non_empty_idx = None
    
    for idx, p in enumerate(paragraphs):
        text = p.text.strip()
        if not text:
            continue
            
        if first_non_empty_idx is None:
            first_non_empty_idx = idx
            
        if p.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            for run in p.runs:
                if run.font.size and run.font.size >= Pt(20):
                    return idx
    
    return first_non_empty_idx


# --------------- 附件块处理 ---------------
def adjust_attachment_block(paragraphs: List[Paragraph], re_attach, re_h1, re_h2) -> None:
    """调整附件块的格式和缩进。
    
    Args:
        paragraphs: 段落列表
        re_attach: 附件匹配正则
        re_h1: 一级标题正则
        re_h2: 二级标题正则
    """
    i = 0
    while i < len(paragraphs):
        p = paragraphs[i]
        text = p.text.strip()
        if re_attach.match(text):
            block_start = i
            block_end = block_start + 1
            while block_end < len(paragraphs):
                next_text = paragraphs[block_end].text.strip()
                if re_h1.match(next_text) or re_h2.match(next_text) or re_attach.match(next_text):
                    break
                block_end += 1
            _align_attachment_with_indent(paragraphs, block_start, block_end)
            i = block_end
        else:
            i += 1


def _align_attachment_with_indent(paragraphs: List[Paragraph], start: int, end: int) -> None:
    """对齐附件块内段落的缩进。
    
    Args:
        paragraphs: 段落列表
        start: 起始索引
        end: 结束索引
    """
    first_p = paragraphs[start]
    strip_leading_spaces(first_p)
    first_text = first_p.text
    prefix_match = RE_ATTACH_PREFIX.match(first_text)
    if not prefix_match:
        for i in range(start, end):
            paragraphs[i].paragraph_format.first_line_indent = pt_from_chars(BODY_SIZE, 2)
        return
    prefix = prefix_match.group(1)
    prefix_len = len(prefix)
    first_p.paragraph_format.first_line_indent = pt_from_chars(BODY_SIZE, 2)
    for i in range(start + 1, end):
        p = paragraphs[i]
        strip_leading_spaces(p)
        text = p.text
        if RE_ATTACH_INDENT.search(text):
            indent_chars = 2 + prefix_len
            p.paragraph_format.first_line_indent = pt_from_chars(BODY_SIZE, indent_chars)
        else:
            p.paragraph_format.first_line_indent = pt_from_chars(BODY_SIZE, 2)


# ---------------------------- 落款处理（支持偏移）-------------------------------
def apply_signature_style(paragraphs: List[Paragraph], signature_style, body_style, doc: DocumentType) -> None:
    """应用落款样式，自动识别并格式化发文单位和日期。
    
    Args:
        paragraphs: 段落列表
        signature_style: 落款样式
        body_style: 正文样式
        doc: Document对象
    """
    max_search = min(15, len(paragraphs))
    start_idx = max(0, len(paragraphs) - max_search)

    date_idx = None
    for i in range(len(paragraphs) - 1, start_idx - 1, -1):
        text = paragraphs[i].text.strip()
        if RE_DATE_PATTERN.search(text):
            date_idx = i
            break
    if date_idx is None:
        return

    p_date = paragraphs[date_idx]
    date_text = p_date.text.strip()
    if not date_text:
        return
    p_date.clear()
    run = p_date.add_run(date_text)
    set_run_font(run, BODY_FONT, BODY_SIZE, set_west=False)
    p_date.paragraph_format.first_line_indent = Pt(0)

    unit_idx = date_idx - 1
    while unit_idx >= start_idx:
        text = paragraphs[unit_idx].text.strip()
        if text and not RE_SIGNATURE_EXCLUDE.match(text):
            break
        unit_idx -= 1
    if unit_idx < 0:
        return

    p_unit = paragraphs[unit_idx]
    unit_text = p_unit.text.strip()
    if not unit_text:
        return
    p_unit.clear()
    run = p_unit.add_run(unit_text)
    set_run_font(run, BODY_FONT, BODY_SIZE, set_west=False)
    p_unit.paragraph_format.first_line_indent = Pt(0)

    clean_date = RE_CLEAN_SPACES.sub('', date_text)
    clean_unit = RE_CLEAN_SPACES.sub('', unit_text)
    date_w = sum(char_width(ch) for ch in clean_date)
    unit_w = sum(char_width(ch) for ch in clean_unit)

    base_indent = 4.0
    unit_indent = base_indent - (unit_w - date_w) / 2.0 + SIGNATURE_OFFSET + 0.7

    logger.debug(f"落款：单位={unit_text!r}(宽{unit_w}) 日期={date_text!r}(宽{date_w}) → 单位右缩进={unit_indent:.2f}字符")

    p_unit.style = signature_style
    p_unit.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_unit.paragraph_format.right_indent = pt_from_chars(BODY_SIZE, unit_indent)

    p_date.style = signature_style
    p_date.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_date.paragraph_format.right_indent = pt_from_chars(BODY_SIZE, base_indent)

    if unit_idx > 0 and paragraphs[unit_idx - 1].text.strip():
        insert_empty_paragraph_before(p_unit, doc, body_style)


# ---------------------------- 二级标题混合格式 -------------------------------
def apply_mixed_h2_format(paragraph: Paragraph, h2_style, body_style) -> None:
    """应用二级标题混合格式（前半部分楷体加粗，后半部分仿宋）。
    
    Args:
        paragraph: 段落对象
        h2_style: 二级标题样式
        body_style: 正文样式
    """
    text = paragraph.text.strip()
    if not text:
        return
    match = RE_H2_SPLIT.search(text)
    if not match:
        paragraph.style = h2_style
        paragraph.paragraph_format.first_line_indent = pt_from_chars(BODY_SIZE, 2)
        for run in paragraph.runs:
            run._element.getparent().remove(run._element)
        new_run = paragraph.add_run(text)
        set_run_font(new_run, "楷体_GB2312", BODY_SIZE, bold=True)
        return
    paragraph.style = body_style
    split_pos = match.end()
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)
    run1 = paragraph.add_run(text[:split_pos])
    set_run_font(run1, "楷体_GB2312", BODY_SIZE, bold=True)
    run2 = paragraph.add_run(text[split_pos:])
    set_run_font(run2, BODY_FONT, BODY_SIZE, bold=False)
    paragraph.paragraph_format.first_line_indent = pt_from_chars(BODY_SIZE, 2)


# ---------------------------- 页码添加函数 ---------------------------------
def add_page_number(paragraph: Paragraph, fmt: str = PAGE_FORMAT) -> None:
    """在段落中添加页码域。
    
    Args:
        paragraph: 段落对象
        fmt: 页码格式模板，使用 {PAGE} 占位符
    """
    parts = fmt.split('{PAGE}')
    if len(parts) == 2:
        left, right = parts
    else:
        left, right = '— ', ' —'
    run_left = paragraph.add_run(left)
    run_left.font.name = 'Times New Roman'
    run_left.font.size = Pt(FOOTER_FONT_SIZE)
    run_field = paragraph.add_run()
    run_field.font.name = 'Times New Roman'
    run_field.font.size = Pt(FOOTER_FONT_SIZE)
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run_field._element.append(fldChar_begin)
    instrText = OxmlElement('w:instrText')
    instrText.text = ' PAGE '
    run_field._element.append(instrText)
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run_field._element.append(fldChar_end)
    run_right = paragraph.add_run(right)
    run_right.font.name = 'Times New Roman'
    run_right.font.size = Pt(FOOTER_FONT_SIZE)


def set_page_margins(doc: DocumentType) -> None:
    """设置页面边距为公文标准。
    
    Args:
        doc: Document对象
    """
    for section in doc.sections:
        section.page_height = PAGE_HEIGHT
        section.page_width = PAGE_WIDTH
        section.top_margin = TOP_MARGIN
        section.bottom_margin = BOTTOM_MARGIN
        section.left_margin = LEFT_MARGIN
        section.right_margin = RIGHT_MARGIN


def create_all_styles(doc: DocumentType) -> dict:
    """创建所有公文样式。
    
    Args:
        doc: Document对象
        
    Returns:
        样式名称到样式对象的字典
    """
    styles = {}
    for style_name, config in STYLE_CONFIG.items():
        style = get_or_create_style(doc, style_name)
        apply_style_config(style, config)
        styles[style_name] = style
    return styles


def set_table_fonts(doc: DocumentType) -> None:
    """设置表格中未指定字体的运行。
    
    Args:
        doc: Document对象
    """
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if not run.font.name:
                            set_run_font(run, BODY_FONT, BODY_SIZE)


def get_unique_backup_path(backup_path: str) -> str:
    """获取唯一的备份文件路径，避免覆盖已有备份。
    
    如果备份文件已存在，则在文件名后添加递增序号。
    例如：backup.txt → backup(1).txt → backup(2).txt
    
    Args:
        backup_path: 原始备份文件路径
        
    Returns:
        唯一的备份文件路径
    """
    if not os.path.exists(backup_path):
        return backup_path
    
    file_dir = os.path.dirname(backup_path)
    file_name = os.path.basename(backup_path)
    name, ext = os.path.splitext(file_name)
    
    counter = 1
    while True:
        new_file_name = f"{name}({counter}){ext}"
        new_path = os.path.join(file_dir, new_file_name)
        if not os.path.exists(new_path):
            return new_path
        counter += 1


def set_footer_page_number(doc: DocumentType) -> None:
    """设置页脚页码。
    
    Args:
        doc: Document对象
    """
    for section in doc.sections:
        footer = section.footer
        if not footer.paragraphs:
            footer.add_paragraph()
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer_para.runs:
            run._element.getparent().remove(run._element)
        add_page_number(footer_para)


# ------------------------------- 主格式化函数 ---------------------------------
def format_official_document(input_path: str, output_path: str, backup_path: Optional[str] = None) -> str:
    """格式化公文文档。
    
    Args:
        input_path: 输入文件路径
        output_path: 输出文件路径
        backup_path: 备份文件路径（可选）
        
    Returns:
        实际使用的备份文件路径，如果未备份则返回None
        
    Raises:
        FileNotFoundError: 输入文件不存在
        OSError: 备份或保存失败
        ValueError: 文档无法打开
    """
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"输入文件不存在：{input_path}")
    
    actual_backup_path = None
    if backup_path:
        try:
            actual_backup_path = get_unique_backup_path(backup_path)
            shutil.copy2(input_path, actual_backup_path)
            logger.info(f"已备份原文件：{actual_backup_path}")
        except OSError as e:
            raise OSError(f"备份文件失败：{e}")
    
    try:
        doc = Document(input_path)
    except Exception as e:
        raise ValueError(f"无法打开文档：{e}")

    set_page_margins(doc)
    styles = create_all_styles(doc)
    
    title_style = styles['公文标题']
    h1_style = styles['一级标题']
    h2_style = styles['二级标题']
    main_receiver_style = styles['主送机关']
    body_style = styles['公文正文']
    signature_style = styles['公文落款']
    attachment_style = styles['公文附件']

    paragraphs = doc.paragraphs
    for p in paragraphs:
        strip_leading_spaces(p)

    title_idx = find_title_index(paragraphs)
    if title_idx is not None:
        paragraphs[title_idx].style = title_style
        for run in paragraphs[title_idx].runs:
            clear_run_font(run)

    main_receiver_idx = None
    start_search = title_idx + 1 if title_idx is not None else 0
    for i in range(start_search, len(paragraphs)):
        text = paragraphs[i].text.strip()
        if not text:
            continue
        if RE_H1.match(text) or RE_H2.match(text) or RE_ATTACH.match(text):
            break
        if RE_MAIN_RECEIVER.match(text):
            main_receiver_idx = i
            break

    if main_receiver_idx is not None and title_idx is not None:
        between = paragraphs[title_idx + 1:main_receiver_idx]
        if all(p.text.strip() != '' for p in between):
            insert_empty_paragraph_before(paragraphs[main_receiver_idx], doc, body_style)
            paragraphs = doc.paragraphs
            main_receiver_idx = next(
                (i for i, p in enumerate(paragraphs) if p._element is paragraphs[main_receiver_idx]._element),
                main_receiver_idx
            )

    adjust_attachment_block(paragraphs, RE_ATTACH, RE_H1, RE_H2)

    attach_indices = [i for i, p in enumerate(paragraphs) if RE_ATTACH.match(p.text.strip())]
    for idx in attach_indices:
        if idx > 0 and paragraphs[idx - 1].text.strip():
            insert_empty_paragraph_before(paragraphs[idx], doc, body_style)
            paragraphs = doc.paragraphs
            break

    in_attachment_block = False
    for p in paragraphs:
        if title_idx is not None and p._element is paragraphs[title_idx]._element:
            continue
        text = p.text.strip()
        if not text:
            p.style = body_style
            for run in p.runs:
                clear_run_font(run)
            continue
        if RE_H1.match(text):
            in_attachment_block = False
            p.style = h1_style
            set_paragraph_font(p, '黑体', BODY_SIZE, bold=False)
        elif RE_H2.match(text):
            in_attachment_block = False
            apply_mixed_h2_format(p, h2_style, body_style)
        elif RE_ATTACH.match(text):
            in_attachment_block = True
            p.style = attachment_style
            set_paragraph_font(p, BODY_FONT, BODY_SIZE)
        elif in_attachment_block:
            p.style = attachment_style
            set_paragraph_font(p, BODY_FONT, BODY_SIZE)
        elif main_receiver_idx is not None and p._element is paragraphs[main_receiver_idx]._element:
            p.style = main_receiver_style
            set_paragraph_font(p, BODY_FONT, BODY_SIZE)
        else:
            p.style = body_style
            p.paragraph_format.first_line_indent = pt_from_chars(BODY_SIZE, 2)
            set_paragraph_font(p, BODY_FONT, BODY_SIZE)

    apply_signature_style(paragraphs, signature_style, body_style, doc)
    set_table_fonts(doc)
    set_footer_page_number(doc)

    try:
        doc.save(output_path)
        logger.info(f"格式化完成：{output_path}")
    except OSError as e:
        if actual_backup_path and os.path.exists(actual_backup_path):
            try:
                shutil.copy2(actual_backup_path, input_path)
                logger.info(f"已从备份恢复：{actual_backup_path}")
            except OSError:
                logger.error(f"恢复备份失败：{actual_backup_path}")
        raise OSError(f"保存文档失败：{e}")
    
    return actual_backup_path


def main():
    """主函数：处理命令行参数并执行格式化。"""
    print("=" * 50)
    print("公文自动格式化工具 v4.0.0")
    print("作者：茗记")
    print("=" * 50)
    print()
    if len(sys.argv) < 2:
        print("使用方法：将 .docx 文件拖到此图标上，或命令行：python 工具名.py 文件1.docx ...")
        print()
        input("按回车键退出...")
        return

    success = 0
    fail = 0
    total = len(sys.argv) - 1
    for i, input_file in enumerate(sys.argv[1:], 1):
        print(f"正在处理第 {i}/{total} 个文件：{os.path.basename(input_file)}")
        if not input_file.lower().endswith('.docx'):
            print("  ❌ 错误：只支持 .docx 格式")
            fail += 1
            continue
        
        file_dir = os.path.dirname(input_file)
        name, ext = os.path.splitext(os.path.basename(input_file))
        backup_file = os.path.join(file_dir, f"{name}-旧{ext}")
        output_file = input_file
        
        try:
            actual_backup = format_official_document(input_file, output_file, backup_file)
            print(f"  ✅ 格式化完成！")
            print(f"     📁 原文件已替换为格式化版本")
            if actual_backup:
                print(f"     💾 备份文件：{os.path.basename(actual_backup)}")
            success += 1
        except FileNotFoundError as e:
            print(f"  ❌ 文件不存在：{str(e)}")
            fail += 1
        except OSError as e:
            print(f"  ❌ 系统错误：{str(e)}")
            fail += 1
        except ValueError as e:
            print(f"  ❌ 文档错误：{str(e)}")
            fail += 1
        except Exception as e:
            logger.exception(f"处理失败：{input_file}")
            print(f"  ❌ 未知错误：{str(e)}")
            fail += 1
    print("=" * 50)
    print(f"处理完成！成功 {success} 个，失败 {fail} 个")
    print("=" * 50)
    input("按回车键退出...")


if __name__ == "__main__":
    main()
