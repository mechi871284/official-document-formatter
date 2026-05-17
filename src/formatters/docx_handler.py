"""
公文格式化工具 - DOCX格式处理器
处理Microsoft Word .docx格式文档
"""

import os
import re
import shutil
import logging
from functools import lru_cache
from typing import Optional, List, Tuple
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
from docx.text.paragraph import Paragraph
from docx.document import Document as DocumentType
from docx.styles.style import _ParagraphStyle

from .base_handler import FormatHandler

logger = logging.getLogger(__name__)

# ------------------------------ 可调参数 ---------------------------------------
BODY_FONT = '仿宋_GB2312'
BODY_SIZE = 16
PAGE_FORMAT = '— PAGE —'
SIGNATURE_OFFSET = 0.0  # 单位左移校正值（正数左移，负数右移），单位：全角字符

# ------------------------------ 落款对齐常量 -----------------------------------
# 落款右缩进基准值（单位：全角字符）
SIGNATURE_BASE_INDENT = 4.0
# 单位与日期对齐的微调值（单位：全角字符）
SIGNATURE_ALIGN_ADJUST = 0.7

# ------------------------------ 样式配置常量 -----------------------------------
STYLE_CONFIG = {
    '公文标题': {
        'font': '方正小标宋简体',
        'size': 22,
        'bold': False,
        'alignment': WD_ALIGN_PARAGRAPH.CENTER,
        'first_line_indent': None,
        'space_before': 0,
        'space_after': 0,
        'line_spacing': 32,
        'set_west': False,
    },
    '一级标题': {
        'font': '黑体',
        'size': BODY_SIZE,
        'bold': False,
        'alignment': None,
        'first_line_indent': 0,
        'space_before': 0,
        'space_after': 0,
        'line_spacing': 28,
        'set_west': False,
    },
    '二级标题': {
        'font': '楷体_GB2312',
        'size': BODY_SIZE,
        'bold': True,
        'alignment': None,
        'first_line_indent': 0,
        'space_before': 0,
        'space_after': 0,
        'line_spacing': 28,
        'set_west': False,
    },
    '主送机关': {
        'font': BODY_FONT,
        'size': BODY_SIZE,
        'bold': False,
        'alignment': WD_ALIGN_PARAGRAPH.LEFT,
        'first_line_indent': 0,
        'space_before': 0,
        'space_after': 0,
        'line_spacing': 28,
        'set_west': False,
    },
    '公文正文': {
        'font': BODY_FONT,
        'size': BODY_SIZE,
        'bold': False,
        'alignment': None,
        'first_line_indent': 2,
        'space_before': 0,
        'space_after': 0,
        'line_spacing': 28,
        'set_west': False,
    },
    '公文落款': {
        'font': BODY_FONT,
        'size': BODY_SIZE,
        'bold': False,
        'alignment': WD_ALIGN_PARAGRAPH.RIGHT,
        'first_line_indent': 0,
        'right_indent': 0,
        'space_before': 0,
        'space_after': 0,
        'line_spacing': 28,
        'set_west': False,
    },
    '公文附件': {
        'font': BODY_FONT,
        'size': BODY_SIZE,
        'bold': False,
        'alignment': None,
        'first_line_indent': 0,
        'space_before': 0,
        'space_after': 0,
        'line_spacing': 28,
        'set_west': False,
    },
}

# ------------------------------ 页面配置常量 -----------------------------------
PAGE_HEIGHT = Cm(29.7)
PAGE_WIDTH = Cm(21.0)
TOP_MARGIN = Cm(3.7)
BOTTOM_MARGIN = Cm(3.5)
LEFT_MARGIN = Cm(2.8)
RIGHT_MARGIN = Cm(2.6)

# ------------------------------ 行距配置常量 -----------------------------------
LINE_SPACING_TITLE = 32
LINE_SPACING_OTHER = 28
FOOTER_FONT_SIZE = 14

# ------------------------------ 正则表达式预编译 -------------------------------
# 一级标题：支持"一、" "1、" "1."等多种序号格式
RE_H1 = re.compile(r'^(?:[一二三四五六七八九十]+[、.])|(?:\d{1,2}[、.])')
# 二级标题：支持"（一）" "(一)" "1)"等多种格式
RE_H2 = re.compile(r'^(?:（[一二三四五六七八九十]+）)|(?:\([一二三四五六七八九十]+\))|(?:\d{1,2}\))')
# 附件：支持"附件" "附件1" "附件：" "附件1："及全角冒号
RE_ATTACH = re.compile(r'^附件\d*\s*[：:]?')
# 主送机关：更宽松的匹配模式，支持冒号结尾或特定模式
RE_MAIN_RECEIVER = re.compile(
    r'^(?:各|各有关|本|本司|本局|各部|各省|自治区|直辖市).*[：:]\s*$'
    r'|^(?!.*(?:注明|备注|说明|解释|参考|来源|联系|电话|传真|电子邮箱|附件|抄送|主题词|印发日期))[^。；，！？…]{3,30}[：:]\s*$'
)
# 日期模式：增强支持各种日期格式变体
RE_DATE_PATTERN = re.compile(
    r'(\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日)'
    r'|((?:二[零〇一二三四五六七八九十]+)\s*年\s*.*?\s*月\s*.*?\s*日)'
    r'|(\d{4}[-/]\d{1,2}[-/]\d{1,2})'
)
# 附件前缀：匹配到冒号为止（包括空格）
RE_ATTACH_PREFIX = re.compile(r'^(附件\d*\s*[：:])')
# 附件缩进识别：增强支持"1." "1、" "1)"等
RE_ATTACH_INDENT = re.compile(r'^\s*\d+[.、)）]')
# 落款排除
RE_SIGNATURE_EXCLUDE = re.compile(r'附件|抄送|主题词|（此页无正文）|\d{4}年')
# H2分隔符
RE_H2_SPLIT = re.compile(r'[。；]')
# 空白清理（包括全角空格、不间断空格）
RE_CLEAN_SPACES = re.compile(r'[\s\u3000\u00A0\u2000-\u200B]+')


@lru_cache(maxsize=512)
def is_fullwidth(ch: str) -> bool:
    """判断字符是否为全角字符。"""
    code = ord(ch)
    if 0x4E00 <= code <= 0x9FFF:
        return True
    elif 0x3400 <= code <= 0x4DBF:
        return True
    elif 0xFF00 < code <= 0xFFEF:
        return True
    elif code in (0x3000, 0x3001, 0x3002, 0x300A, 0x300B, 0x300C, 0x300D,
                  0x300E, 0x300F, 0x3010, 0x3011, 0x2018, 0x2019,
                  0x201C, 0x201D, 0x2013, 0x2014, 0xFF08, 0xFF09):
        return True
    elif 0xFE30 <= code <= 0xFE4F:
        return True
    return False


def char_width(ch: str) -> float:
    """计算字符宽度（以全角字符为单位）。"""
    if is_fullwidth(ch):
        return 1.0
    if ch.isascii():
        return 0.5
    return 1.0


def get_or_create_style(doc: DocumentType, name: str, style_type=WD_STYLE_TYPE.PARAGRAPH) -> _ParagraphStyle:
    """获取或创建文档样式。"""
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, style_type)


def apply_style_config(style: _ParagraphStyle, config: dict) -> None:
    """根据配置字典应用样式设置。"""
    style.font.name = config['font']
    style.font.size = Pt(config['size'])
    style.font.bold = config['bold']
    if config['alignment'] is not None:
        style.paragraph_format.alignment = config['alignment']
    if config['first_line_indent'] is not None:
        if config['first_line_indent'] == 0:
            style.paragraph_format.first_line_indent = Pt(0)
        else:
            style.paragraph_format.first_line_indent = pt_from_chars(BODY_SIZE, config['first_line_indent'])
    if 'right_indent' in config:
        style.paragraph_format.right_indent = Pt(config['right_indent'])
    style.paragraph_format.space_before = Pt(config['space_before'])
    style.paragraph_format.space_after = Pt(config['space_after'])
    style.paragraph_format.line_spacing = Pt(config['line_spacing'])
    rPr = style.element.get_or_add_rPr()
    safe_set_font(rPr, config['font'], config.get('set_west', True))


def pt_from_chars(font_size_pt: float, chars: float = 2.0) -> Pt:
    """将字符数转换为磅值。"""
    return Pt(font_size_pt * chars)


def safe_set_font(rPr, font_name: str, set_west: bool = True) -> None:
    """安全设置字体，确保XML元素存在。
    
    Args:
        rPr: Run properties XML元素
        font_name: 字体名称
        set_west: 是否同时设置西文字体（ascii, hAnsi）
    """
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def set_run_font(run, font_name: str, font_size: float, bold: bool = False, set_west: bool = True) -> None:
    """设置文本运行的字体。"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    safe_set_font(rPr, font_name, set_west)


def set_paragraph_font(paragraph: Paragraph, font_name: str, font_size: float, bold: bool = False) -> None:
    """设置段落中所有运行的字体。"""
    for run in paragraph.runs:
        set_run_font(run, font_name, font_size, bold)


def clear_run_font(run) -> None:
    """清除运行的字体格式。"""
    run.font.name = None
    run.font.size = None
    run.font.bold = None
    run.font.italic = None
    rPr = run._element.find(qn('w:rPr'))
    if rPr is not None:
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is not None:
            for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
                if qn(attr) in rFonts.attrib:
                    del rFonts.attrib[qn(attr)]
            if not rFonts.attrib:
                rPr.remove(rFonts)


def strip_leading_spaces(paragraph: Paragraph) -> None:
    """清除段落开头的所有空白字符。"""
    runs = paragraph.runs
    if not runs:
        return
    full_text = ''.join(run.text for run in runs)
    stripped = full_text.lstrip()
    if stripped == full_text:
        return
    for run in runs:
        run._element.getparent().remove(run._element)
    new_run = paragraph.add_run(stripped)
    set_run_font(new_run, BODY_FONT, BODY_SIZE)


def _clear_paragraph_runs(paragraph: Paragraph) -> None:
    """清除段落中所有运行的直接字体格式，确保样式生效。"""
    for run in paragraph.runs:
        rPr = run._element.find(qn('w:rPr'))
        if rPr is not None:
            run._element.remove(rPr)


def insert_empty_paragraph_before(target_paragraph: Paragraph, doc: DocumentType, body_style) -> None:
    """在目标段落前插入空段落。"""
    empty_p = OxmlElement('w:p')
    target_paragraph._element.addprevious(empty_p)
    for p in doc.paragraphs:
        if p._element is empty_p:
            p.style = body_style
            break


def find_title_index(paragraphs: List[Paragraph]) -> Optional[int]:
    """查找标题段落的索引（单次遍历优化）。"""
    first_non_empty_idx = None
    for idx, p in enumerate(paragraphs):
        text = p.text.strip()
        if not text:
            continue
        if first_non_empty_idx is None:
            first_non_empty_idx = idx
        if p.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            for run in p.runs:
                if run.font.size and run.font.size >= Pt(20):
                    return idx
    return first_non_empty_idx


def adjust_attachment_block(paragraphs: List[Paragraph], re_attach, re_h1, re_h2) -> None:
    """调整附件块的格式和缩进。
    
    公文附件格式规范：
    - 单个附件：附件：关于XXX的通知（首行缩进2字符）
    - 多个附件：附件：（首行缩进2字符）
              1.关于XXX的通知（前端对齐）
              2.关于YYY的说明（前端对齐）
    """
    attach_blocks = []
    i = 0
    while i < len(paragraphs):
        p = paragraphs[i]
        text = p.text.strip()
        if re_attach.match(text):
            block_start = i
            block_end = block_start + 1
            while block_end < len(paragraphs):
                next_text = paragraphs[block_end].text.strip()
                if re_h1.match(next_text) or re_h2.match(next_text) or re_attach.match(next_text):
                    break
                block_end += 1
            attach_blocks.append((block_start, block_end))
            i = block_end
        else:
            i += 1
    
    for start, end in attach_blocks:
        _format_attachment_block(paragraphs, start, end)


def _format_attachment_block(paragraphs: List[Paragraph], start: int, end: int) -> None:
    """格式化单个附件块。
    
    Args:
        paragraphs: 段落列表
        start: 附件块起始索引
        end: 附件块结束索引
    """
    first_p = paragraphs[start]
    strip_leading_spaces(first_p)
    first_text = first_p.text
    
    # 检查是否为多附件格式（后续段落以数字序号开头）
    is_multi = False
    if end - start > 1:
        for i in range(start + 1, end):
            if RE_ATTACH_INDENT.search(paragraphs[i].text.strip()):
                is_multi = True
                break
    
    if is_multi:
        # 多附件格式：引导行缩进2字符，序号行缩进与引导行前缀后内容对齐
        first_p.paragraph_format.first_line_indent = pt_from_chars(BODY_SIZE, 2)
        prefix_match = RE_ATTACH_PREFIX.match(first_text)
        prefix_len = len(prefix_match.group(1)) if prefix_match else 0
        indent_chars = 2 + prefix_len
        for i in range(start + 1, end):
            p = paragraphs[i]
            strip_leading_spaces(p)
            p.paragraph_format.first_line_indent = pt_from_chars(BODY_SIZE, indent_chars)
            for run in p.runs:
                set_run_font(run, BODY_FONT, BODY_SIZE, bold=False)
    else:
        # 单附件格式：首行缩进2字符
        for i in range(start, end):
            paragraphs[i].paragraph_format.first_line_indent = pt_from_chars(BODY_SIZE, 2)
            for run in paragraphs[i].runs:
                set_run_font(run, BODY_FONT, BODY_SIZE, bold=False)


def apply_signature_style(paragraphs: List[Paragraph], signature_style, body_style, doc: DocumentType) -> None:
    """应用落款样式，自动识别并格式化发文单位和日期。"""
    max_search = min(15, len(paragraphs))
    start_idx = max(0, len(paragraphs) - max_search)
    date_idx = None
    for i in range(len(paragraphs) - 1, start_idx - 1, -1):
        text = paragraphs[i].text.strip()
        if RE_DATE_PATTERN.search(text):
            date_idx = i
            break
    if date_idx is None:
        return
    p_date = paragraphs[date_idx]
    date_text = p_date.text.strip()
    if not date_text:
        return
    p_date.clear()
    run = p_date.add_run(date_text)
    set_run_font(run, BODY_FONT, BODY_SIZE, set_west=False)
    p_date.paragraph_format.first_line_indent = Pt(0)
    unit_idx = date_idx - 1
    while unit_idx >= start_idx:
        text = paragraphs[unit_idx].text.strip()
        if text and not RE_SIGNATURE_EXCLUDE.match(text):
            break
        unit_idx -= 1
    if unit_idx < 0:
        return
    p_unit = paragraphs[unit_idx]
    unit_text = p_unit.text.strip()
    if not unit_text:
        return
    p_unit.clear()
    run = p_unit.add_run(unit_text)
    set_run_font(run, BODY_FONT, BODY_SIZE, set_west=False)
    p_unit.paragraph_format.first_line_indent = Pt(0)
    clean_date = RE_CLEAN_SPACES.sub('', date_text)
    clean_unit = RE_CLEAN_SPACES.sub('', unit_text)
    date_w = sum(char_width(ch) for ch in clean_date)
    unit_w = sum(char_width(ch) for ch in clean_unit)
    unit_indent = SIGNATURE_BASE_INDENT - (unit_w - date_w) / 2.0 + SIGNATURE_OFFSET + SIGNATURE_ALIGN_ADJUST
    logger.debug(f"落款：单位={unit_text!r}(宽{unit_w}) 日期={date_text!r}(宽{date_w}) → 单位右缩进={unit_indent:.2f}字符")
    p_unit.style = signature_style
    p_unit.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_unit.paragraph_format.right_indent = pt_from_chars(BODY_SIZE, unit_indent)
    p_date.style = signature_style
    p_date.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_date.paragraph_format.right_indent = pt_from_chars(BODY_SIZE, SIGNATURE_BASE_INDENT)
    if unit_idx > 0 and paragraphs[unit_idx - 1].text.strip():
        insert_empty_paragraph_before(p_unit, doc, body_style)


def apply_mixed_h2_format(paragraph: Paragraph, h2_style, body_style) -> None:
    """应用二级标题混合格式（前半部分楷体加粗，后半部分仿宋）。"""
    text = paragraph.text.strip()
    if not text:
        return
    match = RE_H2_SPLIT.search(text)
    if not match:
        paragraph.style = h2_style
        paragraph.paragraph_format.first_line_indent = pt_from_chars(BODY_SIZE, 2)
        for run in paragraph.runs:
            run._element.getparent().remove(run._element)
        new_run = paragraph.add_run(text)
        set_run_font(new_run, "楷体_GB2312", BODY_SIZE, bold=True)
        return
    paragraph.style = body_style
    split_pos = match.end()
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)
    run1 = paragraph.add_run(text[:split_pos])
    set_run_font(run1, "楷体_GB2312", BODY_SIZE, bold=True)
    run2 = paragraph.add_run(text[split_pos:])
    set_run_font(run2, BODY_FONT, BODY_SIZE, bold=False)
    paragraph.paragraph_format.first_line_indent = pt_from_chars(BODY_SIZE, 2)


def add_page_number(paragraph: Paragraph, fmt: str = PAGE_FORMAT) -> None:
    """在段落中添加页码域。"""
    parts = fmt.split('{PAGE}')
    if len(parts) == 2:
        left, right = parts
    else:
        left, right = '— ', ' —'
    run_left = paragraph.add_run(left)
    run_left.font.name = BODY_FONT
    run_left.font.size = Pt(FOOTER_FONT_SIZE)
    run_field = paragraph.add_run()
    run_field.font.name = BODY_FONT
    run_field.font.size = Pt(FOOTER_FONT_SIZE)
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run_field._element.append(fldChar_begin)
    instrText = OxmlElement('w:instrText')
    instrText.text = ' PAGE '
    run_field._element.append(instrText)
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run_field._element.append(fldChar_end)
    run_right = paragraph.add_run(right)
    run_right.font.name = BODY_FONT
    run_right.font.size = Pt(FOOTER_FONT_SIZE)


def set_page_margins(doc: DocumentType) -> None:
    """设置页面边距为公文标准。"""
    for section in doc.sections:
        section.page_height = PAGE_HEIGHT
        section.page_width = PAGE_WIDTH
        section.top_margin = TOP_MARGIN
        section.bottom_margin = BOTTOM_MARGIN
        section.left_margin = LEFT_MARGIN
        section.right_margin = RIGHT_MARGIN


def create_all_styles(doc: DocumentType) -> dict:
    """创建所有公文样式。"""
    styles = {}
    for style_name, config in STYLE_CONFIG.items():
        style = get_or_create_style(doc, style_name)
        apply_style_config(style, config)
        styles[style_name] = style
    return styles


def set_table_fonts(doc: DocumentType) -> None:
    """设置表格中未指定字体的运行。"""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if not run.font.name:
                            set_run_font(run, BODY_FONT, BODY_SIZE)


def get_unique_backup_path(backup_path: str) -> str:
    """获取唯一的备份文件路径，避免覆盖已有备份。"""
    if not os.path.exists(backup_path):
        return backup_path
    file_dir = os.path.dirname(backup_path)
    file_name, file_ext = os.path.splitext(os.path.basename(backup_path))
    counter = 1
    while True:
        new_backup_path = os.path.join(file_dir, f"{file_name}({counter}){file_ext}")
        if not os.path.exists(new_backup_path):
            return new_backup_path
        counter += 1


def set_footer_page_number(doc: DocumentType) -> None:
    """设置页脚页码。"""
    for section in doc.sections:
        footer = section.footer
        if footer.is_linked_to_previous:
            footer.is_linked_to_previous = False
        if not footer.paragraphs:
            footer.add_paragraph()
        footer_para = footer.paragraphs[0]
        footer_para.clear()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number(footer_para)


def format_docx_content(doc: DocumentType) -> None:
    """格式化DOCX文档内容。"""
    styles = create_all_styles(doc)
    title_style = styles['公文标题']
    h1_style = styles['一级标题']
    h2_style = styles['二级标题']
    receiver_style = styles['主送机关']
    body_style = styles['公文正文']
    signature_style = styles['公文落款']
    attach_style = styles['公文附件']
    set_page_margins(doc)
    set_table_fonts(doc)
    paragraphs = doc.paragraphs
    title_idx = find_title_index(paragraphs)
    if title_idx is not None:
        paragraphs[title_idx].style = title_style
        paragraphs[title_idx].paragraph_format.first_line_indent = None
    main_receiver_idx = None
    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        if text and RE_MAIN_RECEIVER.match(text):
            main_receiver_idx = i
            break
    if main_receiver_idx is not None:
        p = paragraphs[main_receiver_idx]
        p.style = receiver_style
        p.paragraph_format.first_line_indent = Pt(0)
    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        if not text:
            p.style = body_style
            _clear_paragraph_runs(p)
            continue
        if i == title_idx or i == main_receiver_idx:
            continue
        if RE_H1.match(text):
            p.style = h1_style
            p.paragraph_format.first_line_indent = Pt(0)
        elif RE_H2.match(text):
            apply_mixed_h2_format(p, h2_style, body_style)
        elif RE_ATTACH.match(text):
            p.style = attach_style
            p.paragraph_format.first_line_indent = pt_from_chars(BODY_SIZE, 2)
        else:
            p.style = body_style
    adjust_attachment_block(paragraphs, RE_ATTACH, RE_H1, RE_H2)
    apply_signature_style(paragraphs, signature_style, body_style, doc)
    set_footer_page_number(doc)


class DocxFormatHandler(FormatHandler):
    """DOCX格式处理器。"""
    
    @property
    def supported_extensions(self) -> list:
        return ['.docx']
    
    def can_handle(self, file_path: str) -> bool:
        return file_path.lower().endswith('.docx')
    
    def validate_format(self, file_path: str) -> Tuple[bool, str]:
        """验证DOCX文件格式。"""
        if not os.path.exists(file_path):
            return False, f"文件不存在：{file_path}"
        try:
            doc = Document(file_path)
            return True, "DOCX格式验证通过"
        except Exception as e:
            return False, f"DOCX格式验证失败：{str(e)}"
    
    def format_document(self, input_path: str, output_path: str,
                       backup_path: Optional[str] = None) -> Optional[str]:
        """格式化DOCX文档。"""
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"输入文件不存在：{input_path}")
        
        actual_backup_path = None
        if backup_path:
            actual_backup_path = get_unique_backup_path(backup_path)
            try:
                shutil.copy2(input_path, actual_backup_path)
                logger.info(f"已备份原文件：{actual_backup_path}")
            except OSError as e:
                raise OSError(f"备份文件失败：{e}")
        
        try:
            doc = Document(input_path)
        except Exception as e:
            raise ValueError(f"无法打开文档：{e}")
        
        try:
            format_docx_content(doc)
        except Exception as e:
            raise ValueError(f"文档格式化失败：{e}")
        
        try:
            doc.save(output_path)
            logger.info(f"格式化完成：{output_path}")
        except OSError as e:
            if actual_backup_path and os.path.exists(actual_backup_path):
                try:
                    shutil.copy2(actual_backup_path, input_path)
                    logger.info(f"已从备份恢复：{actual_backup_path}")
                except OSError:
                    logger.error(f"恢复备份失败：{actual_backup_path}")
            raise OSError(f"保存文档失败：{e}")
        
        return actual_backup_path
